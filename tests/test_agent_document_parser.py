"""Tests for app/agents/document_parser.py — sync, needs real fixture files."""
import pytest
from docx import Document

from app.agents.document_parser import parse_document, parse_docx


@pytest.fixture
def docx_path(tmp_path):
    """Create a minimal valid DOCX file."""
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("jane@example.com")
    doc.add_paragraph("Senior Engineer at Acme Corp")
    path = str(tmp_path / "test.docx")
    doc.save(path)
    return path


@pytest.fixture
def docx_with_table(tmp_path):
    """Create a DOCX with a table."""
    doc = Document()
    doc.add_paragraph("Resume")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Level"
    table.cell(1, 0).text = "Python"
    table.cell(1, 1).text = "Expert"
    path = str(tmp_path / "table.docx")
    doc.save(path)
    return path


class TestParseDocx:
    def test_basic_docx(self, docx_path):
        result = parse_docx(docx_path)
        assert result.status == "SUCCESS"
        assert result.extraction_method == "DOCX_PARSER"
        assert len(result.raw_blocks) >= 3
        contents = [b.content for b in result.raw_blocks]
        assert "Jane Doe" in contents

    def test_docx_with_table(self, docx_with_table):
        result = parse_docx(docx_with_table)
        assert result.status == "SUCCESS"
        table_blocks = [b for b in result.raw_blocks if b.type == "TABLE"]
        assert len(table_blocks) >= 1

    def test_empty_docx(self, tmp_path):
        doc = Document()
        path = str(tmp_path / "empty.docx")
        doc.save(path)
        result = parse_docx(path)
        assert result.status == "SUCCESS"
        assert len(result.raw_blocks) == 0


class TestParseDocument:
    def test_docx_mime_type(self, docx_path):
        result = parse_document(
            docx_path,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        assert result.status == "SUCCESS"

    def test_unsupported_mime_type(self, docx_path):
        result = parse_document(docx_path, "image/png")
        assert result.status == "FAILED"
        assert result.extraction_method == "NONE"
