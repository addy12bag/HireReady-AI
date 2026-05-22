import logging

import pdfplumber
from docx import Document

from app.schemas.blocks import RawBlock, BoundingBox, DocumentMetadata, ExtractionResult

logger = logging.getLogger(__name__)


def classify_pdf(path: str) -> str:
    """Classify PDF as DIGITAL, SCANNED, or MIXED based on text density."""
    with pdfplumber.open(path) as pdf:
        total_chars = 0
        total_area = 0
        for page in pdf.pages:
            text = page.extract_text() or ""
            total_chars += len(text.strip())
            total_area += page.width * page.height

    if total_area == 0:
        return "SCANNED"

    char_density = total_chars / (total_area / 1000)

    if char_density < 0.5:
        return "SCANNED"
    elif char_density < 2:
        return "MIXED"
    return "DIGITAL"


def parse_pdf(path: str) -> ExtractionResult:
    """Extract text blocks from a PDF file."""
    pdf_type = classify_pdf(path)
    if pdf_type == "SCANNED":
        return ExtractionResult(
            status="FAILED",
            extraction_method="PDFPLUMBER",
            raw_blocks=[],
            document_metadata=DocumentMetadata(),
        )

    blocks: list[RawBlock] = []
    with pdfplumber.open(path) as pdf:
        page_count = len(pdf.pages)
        for page_num, page in enumerate(pdf.pages, 1):
            words = page.extract_words(keep_blank_chars=True, use_text_flow=True)
            for i, word in enumerate(words):
                blocks.append(
                    RawBlock(
                        block_id=f"p{page_num}_w{i}",
                        type="TEXT",
                        content=word.get("text", ""),
                        bbox=BoundingBox(
                            x0=word.get("x0", 0),
                            y0=word.get("top", 0),
                            x1=word.get("x1", 0),
                            y1=word.get("bottom", 0),
                        ),
                        page=page_num,
                        confidence=0.95,
                    )
                )

    return ExtractionResult(
        status="SUCCESS",
        extraction_method="PDFPLUMBER",
        raw_blocks=blocks,
        document_metadata=DocumentMetadata(page_count=page_count),
    )


def parse_docx(path: str) -> ExtractionResult:
    """Extract text blocks from a DOCX file."""
    doc = Document(path)
    blocks: list[RawBlock] = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            blocks.append(
                RawBlock(
                    block_id=f"para_{i}",
                    type="TEXT",
                    content=text,
                    page=1,
                    confidence=1.0,
                )
            )

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                blocks.append(
                    RawBlock(
                        block_id=f"table_{table_idx}_row_{row_idx}",
                        type="TABLE",
                        content=row_text,
                        page=1,
                        confidence=0.9,
                    )
                )

    return ExtractionResult(
        status="SUCCESS",
        extraction_method="DOCX_PARSER",
        raw_blocks=blocks,
        document_metadata=DocumentMetadata(page_count=1),
    )


def parse_document(path: str, mime_type: str) -> ExtractionResult:
    """Route to appropriate parser based on MIME type."""
    if mime_type == "application/pdf":
        return parse_pdf(path)
    elif mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return parse_docx(path)
    else:
        logger.warning(f"Unsupported MIME type: {mime_type}")
        return ExtractionResult(
            status="FAILED",
            extraction_method="NONE",
            raw_blocks=[],
            document_metadata=DocumentMetadata(),
        )
