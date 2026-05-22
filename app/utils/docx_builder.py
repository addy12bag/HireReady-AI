"""Build ATS-friendly DOCX resumes matching a clean LaTeX-inspired layout."""
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from app.schemas.resume import ResumeSchema


def _section_rule(doc: Document):
    """Add a thin horizontal rule below a section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): '000000',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_section_heading(doc: Document, text: str):
    """Section heading with rule — matches LaTeX \\section{} style."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 0, 0)
    _section_rule(doc)


def _add_tabular_line(doc: Document, left_text: str, right_text: str,
                       left_bold: bool = True, left_italic: bool = False,
                       right_italic: bool = False, font_size: float = 10):
    """Two-column tabular row (like LaTeX tabular* with \\fill)."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Remove table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto',
        })
        borders.append(el)
    tblPr.append(borders)
    tbl.append(tblPr)

    # Left cell
    left_cell = table.cell(0, 0)
    p = left_cell.paragraphs[0]
    run = p.add_run(left_text)
    run.font.size = Pt(font_size)
    run.font.name = "Calibri"
    run.bold = left_bold
    run.italic = left_italic

    # Right cell
    right_cell = table.cell(0, 1)
    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(right_text)
    run.font.size = Pt(font_size)
    run.font.name = "Calibri"
    run.italic = right_italic

    # Set column widths
    table.columns[0].width = Inches(4.2)
    table.columns[1].width = Inches(2.3)

    # Reduce spacing
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(0)


def _add_bullet(doc: Document, text: str, font_size: float = 10):
    """Add a bullet point with • marker."""
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25)


def build_docx(resume: ResumeSchema, output_path: str | None = None) -> bytes:
    """Build an ATS-safe DOCX matching a clean LaTeX-inspired layout."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # Reduce default margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # ── Header: Name + Contact ──
    contact = resume.contact
    name_text = contact.name or "Resume"

    # Two-column header: Name left, email/mobile right
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Remove borders
    tbl = header_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto',
        })
        borders.append(el)
    tblPr.append(borders)
    tbl.append(tblPr)

    # Left: Name
    left_cell = header_table.cell(0, 0)
    p = left_cell.paragraphs[0]
    run = p.add_run(name_text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Calibri"

    # Right: Email + Phone
    right_cell = header_table.cell(0, 1)
    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    contact_lines = []
    if contact.email:
        contact_lines.append(f"Email: {contact.email}")
    if contact.phone:
        contact_lines.append(f"Mobile: {contact.phone}")
    run = p.add_run("\n".join(contact_lines))
    run.font.size = Pt(10)
    run.font.name = "Calibri"

    header_table.columns[0].width = Inches(4.2)
    header_table.columns[1].width = Inches(2.3)

    # Portfolio/GitHub line below
    link_parts = []
    if contact.github:
        link_parts.append(f"Github: {contact.github}")
    if contact.linkedin:
        link_parts.append(f"LinkedIn: {contact.linkedin}")
    loc_parts = []
    if contact.location and contact.location.city:
        loc_parts.append(contact.location.city)
    if contact.location and contact.location.country:
        loc_parts.append(contact.location.country)

    if link_parts:
        link_table = doc.add_table(rows=1, cols=2)
        link_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl = link_table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
        borders = tblPr.makeelement(qn('w:tblBorders'), {})
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = borders.makeelement(qn(f'w:{edge}'), {
                qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto',
            })
            borders.append(el)
        tblPr.append(borders)
        tbl.append(tblPr)

        left_cell = link_table.cell(0, 0)
        p = left_cell.paragraphs[0]
        run = p.add_run(link_parts[0])
        run.font.size = Pt(10)
        run.font.name = "Calibri"

        if loc_parts:
            right_cell = link_table.cell(0, 1)
            p = right_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(", ".join(loc_parts))
            run.font.size = Pt(10)
            run.font.name = "Calibri"

        link_table.columns[0].width = Inches(4.2)
        link_table.columns[1].width = Inches(2.3)

    # ── Summary ──
    if resume.summary:
        _add_section_heading(doc, "Professional Summary")
        p = doc.add_paragraph(resume.summary)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = "Calibri"

    # ── Education ──
    if resume.education:
        _add_section_heading(doc, "Education")
        for edu in resume.education:
            # Left: Institution, Right: Location (if we had it)
            _add_tabular_line(
                doc,
                edu.institution or "",
                "",  # no location field in education schema
                left_bold=True, left_italic=False,
            )
            # Left: Degree + GPA, Right: Dates
            degree_parts = []
            if edu.degree and edu.field:
                degree_parts.append(f"{edu.degree} - {edu.field};  GPA: {edu.gpa}" if edu.gpa else f"{edu.degree} - {edu.field}")
            elif edu.degree:
                degree_parts.append(edu.degree)

            date_parts = []
            if edu.start_date:
                date_parts.append(edu.start_date)
            if edu.end_date:
                date_parts.append(edu.end_date)

            if degree_parts:
                _add_tabular_line(
                    doc,
                    degree_parts[0],
                    " - ".join(date_parts),
                    left_bold=False, left_italic=True,
                    right_italic=True, font_size=10,
                )

    # ── Skills Summary ──
    skills = resume.skills
    has_skills = any([skills.technical, skills.soft, skills.languages, skills.certifications])
    if has_skills:
        _add_section_heading(doc, "Skills Summary")
        if skills.technical:
            _add_tabular_line(
                doc,
                "Languages/Technologies",
                ", ".join(skills.technical),
                left_bold=True, font_size=10,
            )
        if skills.soft:
            _add_tabular_line(
                doc,
                "Soft Skills",
                ", ".join(skills.soft),
                left_bold=True, font_size=10,
            )
        if skills.languages:
            _add_tabular_line(
                doc,
                "Spoken Languages",
                ", ".join(skills.languages),
                left_bold=True, font_size=10,
            )
        if skills.certifications:
            _add_tabular_line(
                doc,
                "Certifications",
                ", ".join(skills.certifications),
                left_bold=True, font_size=10,
            )

    # ── Experience ──
    if resume.experience:
        _add_section_heading(doc, "Experience")
        for exp in resume.experience:
            # Left: Company, Right: Location
            _add_tabular_line(
                doc,
                exp.company or "",
                exp.location or "",
                left_bold=True, font_size=10,
            )
            # Left: Role, Right: Dates
            date_parts = []
            if exp.start_date:
                date_parts.append(exp.start_date)
            if exp.end_date:
                date_parts.append(exp.end_date)

            _add_tabular_line(
                doc,
                exp.title or "",
                " - ".join(date_parts),
                left_bold=False, left_italic=True,
                right_italic=True, font_size=10,
            )
            # Bullets
            for bullet in exp.bullets:
                _add_bullet(doc, bullet.text, font_size=10)

    # ── Projects ──
    if resume.projects:
        _add_section_heading(doc, "Projects")
        for proj in resume.projects:
            # Project name with tech in parens
            name = proj.name or "Project"
            if proj.technologies:
                name += f" ({', '.join(proj.technologies)})"

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(name)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"

            if proj.description:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Inches(0.25)
                run = p.add_run(proj.description)
                run.font.size = Pt(10)
                run.font.name = "Calibri"

            if proj.url:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Inches(0.25)
                run = p.add_run(f"URL: {proj.url}")
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                run.font.color.rgb = RGBColor(0, 0, 180)

    # ── Certifications & Achievements ──
    if skills.certifications:
        _add_section_heading(doc, "Certifications & Achievements")
        for cert in skills.certifications:
            _add_bullet(doc, cert, font_size=10)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    if output_path:
        with open(output_path, "wb") as f:
            f.write(buffer.getvalue())

    return buffer.getvalue()
